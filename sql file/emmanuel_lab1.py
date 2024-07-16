from docx import Document

# Create a new Document
doc = Document()

# Add a title
doc.add_heading('CS 331 Computer Organization and Architecture', 0)
doc.add_heading('Summer 2024: Homework 1', 1)
doc.add_heading('PERFORMANCE AND METRICS', 2)
doc.add_heading('Total Points: 100', 2)
doc.add_heading('Due Date: 26th May 2024', 2)

# Add Calculation-Based Questions section
doc.add_heading('Calculation-Based Questions', level=1)

# Add Question 3
doc.add_heading('Question 3: Comparing the Performance of M1 and M2', level=2)

doc.add_heading('Part (a): Which computer is faster for each program, and how many times as fast is it?', level=3)
table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Program'
hdr_cells[1].text = 'Time on M1 (seconds)'
hdr_cells[2].text = 'Time on M2 (seconds)'
hdr_cells[3].text = 'Faster Computer'
hdr_cells[4].text = 'Speedup'
row_cells = table.add_row().cells
row_cells[0].text = '1'
row_cells[1].text = '2.0'
row_cells[2].text = '1.5'
row_cells[3].text = 'M2'
row_cells[4].text = '1.33'
row_cells = table.add_row().cells
row_cells[0].text = '2'
row_cells[1].text = '5.0'
row_cells[2].text = '10.0'
row_cells[3].text = 'M1'
row_cells[4].text = '2.0'

doc.add_paragraph('Calculations:')
doc.add_paragraph('Speedup for Program 1 (M2) = 2.0 seconds / 1.5 seconds = 4/3 ≈ 1.33')
doc.add_paragraph('Speedup for Program 2 (M1) = 10.0 seconds / 5.0 seconds = 2')

# Part (b)
doc.add_heading('Part (b): Instruction execution rate (instructions per second) for each computer when running Program 1', level=3)
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Computer'
hdr_cells[1].text = 'Instructions Executed'
hdr_cells[2].text = 'Time (seconds)'
hdr_cells[3].text = 'Instruction Execution Rate (instructions/second)'
row_cells = table.add_row().cells
row_cells[0].text = 'M1'
row_cells[1].text = '5 × 10^9'
row_cells[2].text = '2.0'
row_cells[3].text = '2.5 × 10^9'
row_cells = table.add_row().cells
row_cells[0].text = 'M2'
row_cells[1].text = '6 × 10^9'
row_cells[2].text = '1.5'
row_cells[3].text = '4 × 10^9'

doc.add_paragraph('Calculations:')
doc.add_paragraph('Instruction execution rate on M1 = 5 × 10^9 / 2.0 seconds = 2.5 × 10^9 instructions/second')
doc.add_paragraph('Instruction execution rate on M2 = 6 × 10^9 / 1.5 seconds = 4 × 10^9 instructions/second')

# Part (c)
doc.add_heading('Part (c): CPI for Program 1 on both machines', level=3)
table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Computer'
hdr_cells[1].text = 'Clock Rate (GHz)'
hdr_cells[2].text = 'Time (seconds)'
hdr_cells[3].text = 'Instructions Executed'
hdr_cells[4].text = 'CPI'
row_cells = table.add_row().cells
row_cells[0].text = 'M1'
row_cells[1].text = '3'
row_cells[2].text = '2.0'
row_cells[3].text = '5 × 10^9'
row_cells[4].text = '1.2'
row_cells = table.add_row().cells
row_cells[0].text = 'M2'
row_cells[1].text = '5'
row_cells[2].text = '1.5'
row_cells[3].text = '6 × 10^9'
row_cells[4].text = '1.25'

doc.add_paragraph('Calculations:')
doc.add_paragraph('CPI on M1 = (3 × 10^9 cycles/second × 2.0 seconds) / 5 × 10^9 instructions = 6 × 10^9 / 5 × 10^9 = 1.2')
doc.add_paragraph('CPI on M2 = (5 × 10^9 cycles/second × 1.5 seconds) / 6 × 10^9 instructions = 7.5 × 10^9 / 6 × 10^9 = 1.25')

# Part (d)
doc.add_heading('Part (d): Performance for a workload where Program 1 must be executed 1600 times each hour', level=3)
table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Computer'
hdr_cells[1].text = 'Program 1 Time (seconds)'
hdr_cells[2].text = 'Program 2 Time (seconds)'
hdr_cells[3].text = 'Total Time (seconds)'
hdr_cells[4].text = 'Program 2 Executions'
row_cells = table.add_row().cells
row_cells[0].text = 'M1'
row_cells[1].text = '3200'
row_cells[2].text = '400'
row_cells[3].text = '3600'
row_cells[4].text = '80'
row_cells = table.add_row().cells
row_cells[0].text = 'M2'
row_cells[1].text = '2400'
row_cells[2].text = '1200'
row_cells[3].text = '3600'
row_cells[4].text = '120'

doc.add_paragraph('Calculations:')
doc.add_paragraph('Time for Program 1 on M1 = 2.0 seconds × 1600 = 3200 seconds')
doc.add_paragraph('Time for Program 1 on M2 = 1.5 seconds × 1600 = 2400 seconds')
doc.add_paragraph('Remaining time on M1 for Program 2 = 3600 seconds - 3200 seconds = 400 seconds')
doc.add_paragraph('Remaining time on M2 for Program 2 = 3600 seconds - 2400 seconds = 1200 seconds')
doc.add_paragraph('Program 2 executions on M1 = 400 seconds / 5.0 seconds per execution = 80')
doc.add_paragraph('Program 2 executions on M2 = 1200 seconds / 10.0 seconds per execution = 120')
doc.add_paragraph('Conclusion: M2 is faster for this workload because it allows more executions of Program 2 within the remaining time.')

# Add Question 4
doc.add_heading('Question 4: Running Program P', level=2)

# Part (a)
doc.add_heading('Part (a): CPU execution time for Program P', level=3)
doc.add_paragraph('Given:')
doc.add_paragraph('Instructions: 7.5 × 10^9')
doc.add_paragraph('Clock rate: 5 GHz = 5 × 10^9 cycles/second')
doc.add_paragraph('CPI: 1.2')

doc.add_heading('Calculations:', level=3)
doc.add_paragraph('CPU execution time = (Instructions × CPI) / Clock rate')
doc.add_paragraph('CPU execution time = (7.5 × 10^9 × 1.2) / 5 × 10^9 = (9 × 10^9) / 5 × 10^9 = 1.8 seconds')

# Part (b)
doc.add_heading('Part (b): Percentage of the CPU time Program P received', level=3)
doc.add_paragraph('Given wall time: 3 seconds')

doc.add_heading('Calculations:', level=3)
doc.add_paragraph('Percentage = (CPU execution time / Wall time) × 100')
doc.add_paragraph('Percentage = (1.8 seconds / 3 seconds) × 100 = 60%')

# Add Question 5
doc.add_heading('Question 5: Comparing Different Implementations of the Same Instruction Set', level=2)

# Part (a)
doc.add_heading('Part (a): Using C1 compiler on both M1 and M2', level=3)
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Class'
hdr_cells[1].text = 'CPI on M1'
hdr_cells[2].text = 'CPI on M2'
hdr_cells[3].text = 'C1 Usage'
row_cells = table.add_row().cells

# Add rows to the table for Part (b)
row_cells = table.add_row().cells
row_cells[0].text = 'A'
row_cells[1].text = '2'
row_cells[2].text = '1'
row_cells[3].text = '40%'
row_cells = table.add_row().cells
row_cells[0].text = 'B'
row_cells[1].text = '3'
row_cells[2].text = '2'
row_cells[3].text = '20%'
row_cells = table.add_row().cells
row_cells[0].text = 'C'
row_cells[1].text = '5'
row_cells[2].text = '2'
row_cells[3].text = '40%'

doc.add_paragraph('Calculations:')
doc.add_paragraph('CPI_avg M1 = (2 × 0.40) + (3 × 0.20) + (5 × 0.40) = 0.8 + 0.6 + 2 = 3.4')
doc.add_paragraph('CPI_avg M2 = (1 × 0.40) + (2 × 0.20) + (2 × 0.40) = 0.4 + 0.4 + 0.8 = 1.6')
doc.add_paragraph('Execution time M1 = 3.4 × 1/(6 × 10^9) = 3.4 / 6 × 10^9 = 0.567 × 10^-9')
doc.add_paragraph('Execution time M2 = 1.6 × 1/(3 × 10^9) = 0.533 × 10^-9')
doc.add_paragraph('Speedup = Execution time M1 / Execution time M2 = 0.567 × 10^-9 / 0.533 × 10^-9 ≈ 1.064')

# Part (c)
doc.add_heading('Part (c): If you purchase M1, which compiler would you use?', level=3)
doc.add_paragraph('The compiler that gives the lowest average CPI on M1 is C1 with an average CPI of 3. So, if you purchase M1, you would use the C1 compiler.')

# Part (d)
doc.add_heading('Part (d): If you purchase M2, which compiler would you use?', level=3)
doc.add_paragraph('The compiler that gives the lowest average CPI on M2 is C2 with an average CPI of 1.6. So, if you purchase M2, you would use the C2 compiler.')

# Part (e)
doc.add_heading('Part (e): Which computer and compiler combination gives the best performance and why?', level=3)
doc.add_paragraph('Best performance is achieved with the lowest execution time, which corresponds to the lowest average CPI and highest clock rate.')
doc.add_paragraph('For M1:')
doc.add_paragraph('Execution time with C1 = 1/(2 × 10^9) = 0.5 × 10^-9')
doc.add_paragraph('For M2:')
doc.add_paragraph('Execution time with C2 = 1.6/(3 × 10^9) = 0.533 × 10^-9')
doc.add_paragraph('Despite M1 having a higher CPI, its higher clock rate results in better performance with the C1 compiler.')

# Add Question 6
doc.add_heading('Question 6: Benchmark Speedup', level=2)
doc.add_heading('How much of the initial execution time would floating-point instructions have to account for to show an overall speedup of 3 on this benchmark?', level=3)
doc.add_paragraph('Given:')
doc.add_paragraph('Total initial execution time: 100 seconds')
doc.add_paragraph('Desired speedup: 3')
doc.add_paragraph('Floating-point hardware improvement: 5 times faster')

doc.add_heading('According to Amdahl\'s Law:', level=3)
doc.add_paragraph('Speedup = 1 / ((1 - P) + (P / S))')
doc.add_paragraph('Where:')
doc.add_paragraph('P is the proportion of execution time affected by the improvement.')
doc.add_paragraph('S is the speedup of the improved portion.')

doc.add_heading('For a speedup of 3:', level=3)
doc.add_paragraph('3 = 1 / (1 - P + (P / 5))')
doc.add_paragraph('Solving for P:')
doc.add_paragraph('3 = 1 / (1 - P + P / 5)')
doc.add_paragraph('1 - P + P / 5 = 1 / 3')
doc.add_paragraph('1 - 0.8P = 0.333')
doc.add_paragraph('0.667 = 0.8P')
doc.add_paragraph('P = 0.667 / 0.8 ≈ 0.834')
doc.add_paragraph('Therefore, floating-point instructions must account for approximately 83.4% of the initial execution time.')

# Add Question 7
doc.add_heading('Question 7: Peak Performance and Execution Speedup', level=2)

# Part (a)
doc.add_heading('Part (a): Peak performance of M1 and M2 in instructions per second', level=3)
doc.add_paragraph('Given:')
doc.add_paragraph('M1 clock rate: 4 GHz')
doc.add_paragraph('M2 clock rate: 6 GHz')
doc.add_heading('Peak performance = Clock rate / CPI', level=3)
doc.add_paragraph('For peak performance, use the instruction with the lowest CPI.')
doc.add_heading('For M1:', level=3)
doc.add_paragraph('Peak performance M1 = 4 × 10^9 cycles/second / 1 cycle/instruction = 4 × 10^9 instructions/second')
doc.add_heading('For M2:', level=3)
doc.add_paragraph('Peak performance M2 = 6 × 10^9 cycles/second / 2 cycles/instruction = 3 × 10^9 instructions/second')

# Part (b)
doc.add_heading('Part (b): Execution speedup if the number of instructions is divided equally among the classes, except class A which occurs twice as often', level=3)
doc.add_paragraph('Given the equal distribution with class A twice as frequent:')
doc.add_paragraph('Usage of A = 2/6 = 1/3')
doc.add_paragraph('Usage of B, C, D, and E = 1/6')
doc.add_heading('Average CPI calculation:', level=3)
doc.add_heading('For M1:', level=3)
doc.add_paragraph('CPI_avg M1 = (1 × 1/3) + (2 × 1/6) + (3 × 1/6) + (4 × 1/6) + (3 × 1/6)')
doc.add_paragraph('CPI_avg M1 = 1/3 + 2 = 7/3 ≈ 2.33')

doc.add_heading('For M2:', level=3)
doc.add_paragraph('CPI_avg M2 = (2 × 1/3) + (2 × 1/6) + (2 × 1/6) + (4 × 1/6) + (4 × 1/6)')
doc.add_paragraph('CPI_avg M2 = 2/3 + 2 = 8/3 ≈ 2.67')

doc.add_heading('Execution time:', level=3)
doc.add_heading('For M1:', level=3)
doc.add_paragraph('Execution time M1 = CPI_avg × 1/Clock rate = 2.33 × 1/(4 × 10^9) = 0.583 × 10^-9')
doc.add_heading('For M2:', level=3)
doc.add_paragraph('Execution time M2 = CPI_avg × 1/Clock rate = 2.67 × 1/(6 × 10^9) = 0.445 × 10^-9')

doc.add_heading('Speedup:', level=3)
doc.add_paragraph('Speedup = Execution time M1 / Execution time M2 = 0.583 × 10^-9 / 0.445 × 10^-9 ≈ 1.31')
doc.add_paragraph('So M2 is approximately 1.31 times faster than M1 under the given instruction distribution.')

# Save the document
file_path = "/mnt/data/CS331_Homework_1_Solutions.docx"
doc.save(file_path)
